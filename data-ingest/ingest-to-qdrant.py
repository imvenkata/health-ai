import os
import logging
from typing import List, Dict, Any, Union
from dataclasses import dataclass, field
from pathlib import Path
from dotenv import load_dotenv
import uuid  # Add this import for generating UUIDs
# PDF Parsing
from langchain_community.document_loaders import PyPDFLoader
# DOCX Parsing
from docx import Document as DocxDocument
# Chunking
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument  # Import LangChain Document
# Embedding
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
# Qdrant DB
from qdrant_client import QdrantClient
from qdrant_client.http.models import PointStruct, VectorParams, Distance
# Load environment variables from .env file
load_dotenv()

def get_log_level(level: Union[str, int]) -> int:
    """Convert string log level to logging constant."""
    if isinstance(level, int):
        return level
    
    level_map = {
        'DEBUG': logging.DEBUG,
        'INFO': logging.INFO,
        'WARNING': logging.WARNING,
        'ERROR': logging.ERROR,
        'CRITICAL': logging.CRITICAL
    }
    
    return level_map.get(level.upper(), logging.INFO)

@dataclass
class DocumentIngestionConfig:
    """Configuration class for document ingestion process."""
    input_dir: Path = field(default_factory=lambda: Path(os.getenv("INPUT_DIR", "data-ingest/data/documents")))
    chunk_size: int = int(os.getenv("CHUNK_SIZE", "500"))
    chunk_overlap: int = int(os.getenv("CHUNK_OVERLAP", "50"))
    embedding_model: str = os.getenv("EMBEDDING_MODEL", "BAAI/bge-large-en-v1.5")
    qdrant_mode: str = os.getenv("QDRANT_MODE", "local")
    qdrant_url: str = os.getenv("QDRANT_URL", "http://localhost:6333")
    qdrant_cloud_url: str = os.getenv("QDRANT_CLOUD_URL", "")
    qdrant_api_key: str = os.getenv("QDRANT_API_KEY", "")
    collection_name: str = os.getenv("QDRANT_COLLECTION_NAME", "medical-collections")
    log_level: int = field(default_factory=lambda: get_log_level(os.getenv("LOG_LEVEL", "INFO")))

@dataclass
class DocumentProcessor:
    """Handles document processing and vector storage pipeline."""
    config: DocumentIngestionConfig = field(default_factory=DocumentIngestionConfig)
    logger: logging.Logger = field(init=False)
    
    def __post_init__(self):
        """Initialize logger and validate configuration."""
        logging.basicConfig(level=self.config.log_level, 
                          format='%(asctime)s - %(levelname)s: %(message)s')
        self.logger = logging.getLogger(self.__class__.__name__)
        self._validate_config()
    
    def _validate_config(self):
        """Validate ingestion configuration."""
        if not self.config.input_dir.exists():
            raise ValueError(f"Input directory does not exist: {self.config.input_dir}")
        
        if self.config.chunk_size <= 0 or self.config.chunk_overlap < 0:
            raise ValueError("Invalid chunk size or overlap parameters")
    
    def _select_embedding_model(self) -> Any:
        """Select and initialize embedding model based on configuration."""
        return HuggingFaceEmbedding(
            model_name=self.config.embedding_model,
            trust_remote_code=True
        )
    
    def load_documents(self) -> List[LangChainDocument]:
        """Load PDF and DOCX documents from the specified input directory."""
        documents = []
        pdf_files = list(self.config.input_dir.glob('*.pdf'))
        docx_files = list(self.config.input_dir.glob('*.docx'))
        
        self.logger.info(f"Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files to process")
        
        # Process PDF files
        for pdf_path in pdf_files:
            try:
                loader = PyPDFLoader(str(pdf_path))
                doc_pages = loader.load()
                documents.extend(doc_pages)
                self.logger.info(f"Processed PDF: {pdf_path.name}")
            except Exception as e:
                self.logger.error(f"Error processing {pdf_path.name}: {e}")
        
        # Process DOCX files
        for docx_path in docx_files:
            try:
                doc = DocxDocument(docx_path)
                full_text = "\n".join([para.text for para in doc.paragraphs])
                # Create a LangChain Document object for DOCX content
                documents.append(LangChainDocument(
                    page_content=full_text,
                    metadata={
                        "source": str(docx_path),
                        "page": 0  # DOCX files don't have pages, so we set it to 0
                    }
                ))
                self.logger.info(f"Processed DOCX: {docx_path.name}")
            except Exception as e:
                self.logger.error(f"Error processing {docx_path.name}: {e}")
        
        return documents
    
    def chunk_documents(self, documents: List[LangChainDocument]) -> List[Dict]:
        """Split documents into manageable chunks."""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap
        )
        
        chunks = text_splitter.split_documents(documents)
        
        # Convert LangChain documents to dictionaries
        chunk_dicts = []
        for i, chunk in enumerate(chunks):
            chunk_dicts.append({
                "id": str(uuid.uuid4()),  # Generate a UUID for each chunk
                "text": chunk.page_content,
                "metadata": {
                    "source": chunk.metadata.get("source", "Unknown"),
                    "page": chunk.metadata.get("page", 0)
                }
            })
        
        self.logger.info(f"Created {len(chunk_dicts)} document chunks")
        return chunk_dicts
    
    def generate_embeddings(self, chunks: List[Dict]) -> List[List[float]]:
        """Generate vector embeddings for document chunks."""
        embedding_model = self._select_embedding_model()
        
        # Extract text from chunks for embedding
        texts = [chunk["text"] for chunk in chunks]
        embeddings = [embedding_model.get_text_embedding(text) for text in texts]
        
        self.logger.info(f"Generated {len(embeddings)} embeddings")
        return embeddings
    
    def store_in_qdrant(self, chunks: List[Dict], embeddings: List[List[float]]):
        """Store document chunks and embeddings in Qdrant DB."""
        # Initialize Qdrant client based on mode
        if self.config.qdrant_mode == "local":
            client = QdrantClient(
                url=self.config.qdrant_url,
                api_key=self.config.qdrant_api_key,
                timeout=120  # Increase timeout to handle larger payloads
            )
        else:
            client = QdrantClient(
                url=self.config.qdrant_cloud_url,
                api_key=self.config.qdrant_api_key,
                timeout=120  # Increase timeout to handle larger payloads
            )
        
        # Check if the collection exists
        collection_name = self.config.collection_name
        vector_size = len(embeddings[0]) if embeddings else 768  # Default to BERT-like embedding size
        
        collections = client.get_collections()
        collection_exists = any(collection.name == collection_name for collection in collections.collections)
        
        if not collection_exists:
            try:
                client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(size=vector_size, distance=Distance.COSINE)
                )
                self.logger.info(f"Created Qdrant collection: {collection_name}")
            except Exception as e:
                self.logger.error(f"Failed to create collection '{collection_name}': {e}")
                raise
        else:
            self.logger.info(f"Collection '{collection_name}' already exists. Skipping creation.")
        
        # Prepare points for insertion
        points = [
            PointStruct(
                id=chunk["id"],
                vector=embedding,
                payload={
                    "text": chunk["text"],
                    "metadata": chunk["metadata"]
                }
            )
            for chunk, embedding in zip(chunks, embeddings)
        ]
        
        # Upsert points in batches to avoid timeouts
        batch_size = 100  # Adjust batch size as needed
        for i in range(0, len(points), batch_size):
            batch = points[i:i + batch_size]
            try:
                client.upsert(collection_name=collection_name, points=batch)
                self.logger.info(f"Upserted batch {i // batch_size + 1} with {len(batch)} points")
            except Exception as e:
                self.logger.error(f"Failed to upsert batch {i // batch_size + 1}: {e}")
                raise
    
    def run_ingestion_pipeline(self):
        """Execute the document processing and storage pipeline."""
        try:
            # Load documents (PDFs and DOCX)
            documents = self.load_documents()
            
            if not documents:
                self.logger.warning("No documents found for processing")
                return
            
            # Chunk documents
            chunks = self.chunk_documents(documents)
            
            # Generate embeddings
            embeddings = self.generate_embeddings(chunks)
            
            # Store in Qdrant
            self.store_in_qdrant(chunks, embeddings)
            
            self.logger.info("Document ingestion pipeline completed successfully")
            
        except Exception as e:
            self.logger.error(f"Pipeline execution failed: {e}")
            raise

def main():
    """Main entry point for document ingestion."""
    config = DocumentIngestionConfig()
    processor = DocumentProcessor(config)
    processor.run_ingestion_pipeline()

if __name__ == "__main__":
    main()